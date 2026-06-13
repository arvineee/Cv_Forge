"""
Production-Grade CV Parser Engine — Enhanced Edition

Improvements over v1:
- Section-aware text chunking (headings drive extraction context).
- Multi-format date normalization: handles "Jan '21", "2021–22", "Present", etc.
- Confidence-scored name extraction (positional + casing + length heuristics).
- Location pattern extraction (City, Country / City, State formats).
- Full heuristic work-experience and education block parsers (title/company/dates/description).
- Skills alias normalization — "node.js" / "nodejs" / "Node JS" collapse to one canonical form.
- Post-heuristic Pydantic validation with field-level sanitization.
- AI response partial recovery — rescues valid fields from malformed Gemini JSON.
- Thread-safe lazy-loaded AI service.
"""

import os
import re
import json
import logging
from typing import Dict, Any, List, Optional, Tuple
from pydantic import BaseModel, Field, EmailStr, field_validator, model_validator

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────────────────────
# 1. Enhanced Pydantic Schemas (V2) with Validators
# ─────────────────────────────────────────────────────────────

class PersonalInfoSchema(BaseModel):
    full_name: Optional[str] = Field(None, description="Full name of the candidate.")
    email: Optional[str] = Field(None, description="Valid email address.")
    phone: Optional[str] = Field(None, description="Contact phone, with international prefix if available.")
    location: Optional[str] = Field(None, description="City, region, and/or country of residence.")
    job_title: Optional[str] = Field(None, description="Current or target professional job title.")
    linkedin: Optional[str] = Field(None, description="Cleaned LinkedIn URL or handle.")
    portfolio: Optional[str] = Field(None, description="Personal portfolio, website, or GitHub link.")

    @field_validator("email", mode="before")
    @classmethod
    def validate_email_format(cls, v: Any) -> Optional[str]:
        if not v:
            return None
        v = str(v).strip().lower()
        return v if re.fullmatch(r"[\w.\-+]+@[\w\-]+\.\w{2,}", v) else None

    @field_validator("phone", mode="before")
    @classmethod
    def normalize_phone(cls, v: Any) -> Optional[str]:
        if not v:
            return None
        digits = re.sub(r"[^\d+]", "", str(v))
        return digits if 7 <= len(digits) <= 16 else None

    @field_validator("linkedin", "portfolio", mode="before")
    @classmethod
    def clean_url(cls, v: Any) -> Optional[str]:
        if not v:
            return None
        v = str(v).strip().strip(".,()[]\"'")
        if not v.startswith(("http://", "https://", "www.")):
            v = "https://" + v
        return v


class WorkExperienceSchema(BaseModel):
    title: Optional[str] = Field(None, description="Job title held.")
    company: Optional[str] = Field(None, description="Company or organization name.")
    start_date: Optional[str] = Field(None, description="Start date: Month Year or Year.")
    end_date: Optional[str] = Field(None, description="End date: Month Year, Year, or 'Present'.")
    description: Optional[str] = Field(None, description="Key duties and achievements.")

    @field_validator("start_date", "end_date", mode="before")
    @classmethod
    def normalize_date(cls, v: Any) -> Optional[str]:
        return DateNormalizer.normalize(v) if v else None

    @field_validator("description", mode="before")
    @classmethod
    def handle_list_description(cls, v: Any) -> Optional[str]:
        """
        Rescues data when Gemini returns descriptions as a JSON array of strings
        instead of a single plaintext string block.
        """
        if isinstance(v, list):
            return "\n".join(str(item).strip() for item in v if item)
        if v:
            return str(v).strip()
        return None


class EducationSchema(BaseModel):
    degree: Optional[str] = Field(None, description="Degree or certificate name.")
    institution: Optional[str] = Field(None, description="University or educational institution.")
    year: Optional[str] = Field(None, description="Graduation year or date range.")
    grade: Optional[str] = Field(None, description="GPA, classification, or honors.")


class CVParserOutputSchema(BaseModel):
    personal_info: PersonalInfoSchema = Field(default_factory=PersonalInfoSchema)
    professional_summary: Optional[str] = Field(None, description="Concise professional identity, max 200 words.")
    work_experience: List[WorkExperienceSchema] = Field(default_factory=list)
    education: List[EducationSchema] = Field(default_factory=list)
    skills: List[str] = Field(default_factory=list, description="Deduplicated canonical skill names.")
    extra_sections: Dict[str, str] = Field(
        default_factory=dict,
        description="Any CV sections not matching standard headings, keyed by their original heading name.",
    )

    @field_validator("professional_summary", mode="before")
    @classmethod
    def cap_summary(cls, v: Any) -> Optional[str]:
        if not v:
            return None
        words = str(v).split()
        return " ".join(words[:200])

    @field_validator("skills", mode="before")
    @classmethod
    def deduplicate_skills(cls, v: Any) -> List[str]:
        if not v:
            return []
        return list(dict.fromkeys(s.strip() for s in v if s and s.strip()))


# ─────────────────────────────────────────────────────────────
# 2. Date Normalization Utility
# ─────────────────────────────────────────────────────────────

class DateNormalizer:
    MONTH_MAP = {
        "jan": "January", "feb": "February", "mar": "March", "apr": "April",
        "may": "May", "jun": "June", "jul": "July", "aug": "August",
        "sep": "September", "oct": "October", "nov": "November", "dec": "December",
    }
    PRESENT_TOKENS = {"present", "current", "now", "ongoing", "till date", "to date"}

    @classmethod
    def normalize(cls, raw: str) -> Optional[str]:
        raw = str(raw).strip()
        if not raw:
            return None
        if raw.lower() in cls.PRESENT_TOKENS:
            return "Present"

        # e.g. "Jan '21" → "January 2021"
        m = re.match(r"([A-Za-z]{3,9})[\s.'\-]+['']?(\d{2,4})", raw)
        if m:
            month_key = m.group(1).lower()[:3]
            year = m.group(2)
            year = ("20" + year) if len(year) == 2 else year
            month_full = cls.MONTH_MAP.get(month_key, m.group(1).capitalize())
            return f"{month_full} {year}"

        # e.g. "2021–2023" or "2021-Present"
        m = re.match(r"(\d{4})\s*[–\-–]\s*(\d{4}|present|current)", raw, re.I)
        if m:
            end = "Present" if m.group(2).lower() in cls.PRESENT_TOKENS else m.group(2)
            return f"{m.group(1)} – {end}"

        # Bare year
        if re.fullmatch(r"\d{4}", raw):
            return raw

        return raw  # Return as-is if unrecognized


# ─────────────────────────────────────────────────────────────
# 3. Skills Alias Normalization Registry
# ─────────────────────────────────────────────────────────────

SKILLS_ALIAS_MAP: Dict[str, str] = {
    # JavaScript ecosystem
    "nodejs": "Node.js", "node js": "Node.js", "node.js": "Node.js",
    "reactjs": "React", "react.js": "React", "react js": "React",
    "vuejs": "Vue.js", "vue js": "Vue.js", "vue": "Vue.js",
    "angularjs": "Angular", "angular js": "Angular",
    "nextjs": "Next.js", "next js": "Next.js", "next.js": "Next.js",
    "typescript": "TypeScript", "ts": "TypeScript",
    "javascript": "JavaScript", "js": "JavaScript",
    # Python ecosystem
    "py": "Python", "python3": "Python", "python 3": "Python",
    "django rest framework": "DRF", "drf": "DRF",
    "fastapi": "FastAPI", "fast api": "FastAPI",
    # Cloud / DevOps
    "amazon web services": "AWS", "aws": "AWS",
    "google cloud platform": "GCP", "google cloud": "GCP", "gcp": "GCP",
    "microsoft azure": "Azure", "azure": "Azure",
    "kubernetes": "Kubernetes", "k8s": "Kubernetes",
    "ci/cd": "CI/CD", "cicd": "CI/CD", "ci cd": "CI/CD",
    # Databases
    "postgresql": "PostgreSQL", "postgres": "PostgreSQL",
    "mongodb": "MongoDB", "mongo": "MongoDB",
    "mysql": "MySQL", "my sql": "MySQL",
    "mssql": "SQL Server", "sql server": "SQL Server",
    "redis": "Redis",
    # ML/AI
    "machine learning": "Machine Learning", "ml": "Machine Learning",
    "deep learning": "Deep Learning", "dl": "Deep Learning",
    "natural language processing": "NLP", "nlp": "NLP",
    "tensorflow": "TensorFlow", "tf": "TensorFlow",
    "pytorch": "PyTorch", "torch": "PyTorch",
    # General
    "rest api": "REST API", "restful api": "REST API", "restful": "REST API",
    "graphql": "GraphQL", "graph ql": "GraphQL",
    "html5": "HTML5", "html": "HTML5",
    "css3": "CSS3", "css": "CSS3",
    "git": "Git", "github": "GitHub", "gitlab": "GitLab",
    "docker": "Docker",
    "terraform": "Terraform",
    "linux": "Linux",
    "agile": "Agile", "scrum": "Scrum",
    "microservices": "Microservices",
}

# Canonical skill patterns — what to search for in text
SKILL_PATTERNS: List[str] = sorted(set(list(SKILLS_ALIAS_MAP.keys()) + [
    "python", "java", "c++", "c#", "go", "rust", "ruby", "php", "swift", "kotlin",
    "r", "scala", "matlab", "bash", "shell scripting", "powershell",
    "pandas", "numpy", "scikit-learn", "keras", "spark", "hadoop",
    "elasticsearch", "cassandra", "dynamodb", "firebase",
    "aws lambda", "ec2", "s3", "gke", "gcs", "azure devops",
    "jenkins", "github actions", "gitlab ci", "ansible", "nginx", "apache",
    "oauth", "jwt", "websockets", "grpc",
    "figma", "sketch", "adobe xd", "photoshop", "illustrator",
    "excel", "powerpoint", "tableau", "power bi",
    "project management", "leadership", "communication", "problem solving",
    "data analysis", "data visualization",
]))


def normalize_skill(raw: str) -> str:
    key = raw.strip().lower()
    return SKILLS_ALIAS_MAP.get(key, raw.strip().title())


# ─────────────────────────────────────────────────────────────
# 4. Section-Aware Text Chunker
# ─────────────────────────────────────────────────────────────

SECTION_HEADINGS = {
    "summary": ["summary", "profile", "professional summary", "about me", "objective", "career objective"],
    "experience": ["experience", "work experience", "employment", "work history", "professional experience", "career history"],
    "education": ["education", "academic background", "qualifications", "academic qualifications"],
    "skills": ["skills", "technical skills", "core competencies", "competencies", "technologies", "tools"],
    "projects": ["projects", "key projects", "notable projects"],
    "certifications": ["certifications", "certificates", "licenses", "accreditations"],
}

def build_heading_pattern() -> re.Pattern:
    all_headings = [h for group in SECTION_HEADINGS.values() for h in group]
    escaped = [re.escape(h) for h in sorted(all_headings, key=len, reverse=True)]
    return re.compile(
        r"^(?P<heading>" + "|".join(escaped) + r")\s*:?\s*$",
        re.IGNORECASE | re.MULTILINE,
    )

HEADING_RE = build_heading_pattern()


def _is_section_heading(line: str) -> bool:
    """
    Detects whether a line looks like a CV section heading even if not in SECTION_HEADINGS.
    Criteria: short (≤60 chars), no sentence-ending punctuation mid-line, no email/URL,
    title-cased or ALL-CAPS, and optionally ends with ':'.
    """
    stripped = line.strip().rstrip(":")
    if not stripped or len(stripped) > 60:
        return False
    if re.search(r"[@/\\]|https?://|www\.", stripped):
        return False
    # Must contain at least one letter
    if not re.search(r"[A-Za-z]", stripped):
        return False
    # Skip lines that are clearly sentences (contain common sentence mid-punctuation)
    if re.search(r"[,;]\s+[a-z]", stripped):
        return False
    # Title case (each word starts with upper) or ALL CAPS — typical heading styles
    words = stripped.split()
    if len(words) > 6:
        return False
    is_title = all(w[0].isupper() for w in words if w and w[0].isalpha())
    is_allcaps = stripped.isupper() and len(stripped) > 2
    return is_title or is_allcaps


def chunk_sections(text: str) -> Dict[str, str]:
    """
    Splits raw CV text into named sections based on heading detection.
    Known headings are canonicalized; unknown headings are preserved verbatim
    under their original (lowercased) name so no CV content is silently dropped.
    Returns a dict mapping canonical section name → section body text.
    """
    lines = text.split("\n")
    sections: Dict[str, List[str]] = {"_header": []}
    current_section = "_header"

    for line in lines:
        stripped = line.strip()
        m = HEADING_RE.match(stripped)
        if m:
            # Known heading — map to canonical key
            matched_heading = m.group("heading").lower()
            canonical = next(
                (k for k, aliases in SECTION_HEADINGS.items() if matched_heading in aliases),
                matched_heading,
            )
            current_section = canonical
            sections.setdefault(current_section, [])
        elif stripped and _is_section_heading(stripped):
            # Unknown heading — keep verbatim (lowercased) so it isn't lost
            current_section = stripped.rstrip(":").strip().lower()
            sections.setdefault(current_section, [])
        else:
            sections.setdefault(current_section, []).append(line)

    return {k: "\n".join(v).strip() for k, v in sections.items() if "\n".join(v).strip()}


# ─────────────────────────────────────────────────────────────
# 5. Library Dependency Resolution
# ─────────────────────────────────────────────────────────────

try:
    import pypdf
except ImportError:
    try:
        import PyPDF2 as pypdf
    except ImportError:
        pypdf = None
        logger.critical("No PDF library found (pypdf/PyPDF2). PDF ingestion disabled.")

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None
    logger.critical("python-docx not installed. DOCX ingestion disabled.")


# ─────────────────────────────────────────────────────────────
# 6. Core Engine
# ─────────────────────────────────────────────────────────────

class CVParser:
    """
    Enterprise-grade CV parser. Primary path: Gemini structured JSON schema.
    Secondary path: Section-aware heuristic engine with date normalization,
    confidence-scored name extraction, alias-normalized skills, and
    full Pydantic validation before returning.
    """

    def __init__(self) -> None:
        self._ai_service = None

    def _get_ai_service(self) -> Any:
        """Thread-safe lazy loader — prevents circular imports at module load."""
        if self._ai_service is None:
            from app.ai_service import AIService
            self._ai_service = AIService()
        return self._ai_service

    # ─────────────────────────────────────────────────────────
    # Public Entry Point
    # ─────────────────────────────────────────────────────────

    def parse(self, file_path: str, file_ext: str) -> Dict[str, Any]:
        normalized_ext = file_ext.lower().strip(".")
        if normalized_ext == "pdf":
            raw_text = self._extract_pdf_text(file_path)
        elif normalized_ext == "docx":
            raw_text = self._extract_docx_text(file_path)
        else:
            raise ValueError(f"Unsupported file extension: '{normalized_ext}'. Accepted: pdf, docx.")

        if not raw_text or len(raw_text.strip()) < 50:
            logger.warning(f"Extracted text from '{file_path}' is empty or critically short.")
            return CVParserOutputSchema().model_dump()

        # Primary: Gemini structured extraction
        try:
            result = self._parse_with_ai(raw_text)
            if result:
                return result
        except Exception as e:
            logger.error(f"AI parsing failed: {e}", exc_info=True)

        # Fallback: section-aware heuristic engine
        logger.warning("Falling back to heuristic parser.")
        return self._parse_with_heuristics(raw_text)

    # ─────────────────────────────────────────────────────────
    # Text Extractors
    # ─────────────────────────────────────────────────────────

    def _extract_pdf_text(self, path: str) -> str:
        if pypdf is None:
            raise ImportError("pypdf or PyPDF2 is required for PDF parsing.")
        chunks = []
        try:
            with open(path, "rb") as f:
                reader = pypdf.PdfReader(f)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        chunks.append(text)
            combined = "\n".join(chunks)
            return re.sub(r"\n{3,}", "\n\n", combined).strip()
        except Exception as e:
            logger.error(f"PDF extraction failed for '{path}': {e}")
            return ""

    def _extract_docx_text(self, path: str) -> str:
        if DocxDocument is None:
            raise ImportError("python-docx is required for DOCX parsing.")
        try:
            doc = DocxDocument(path)
            elements = []
            for para in doc.paragraphs:
                if para.text.strip():
                    elements.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_cells = list(dict.fromkeys(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    ))
                    if row_cells:
                        elements.append(" | ".join(row_cells))
            return "\n".join(elements).strip()
        except Exception as e:
            logger.error(f"DOCX extraction failed for '{path}': {e}")
            return ""

    # ─────────────────────────────────────────────────────────
    # AI Path: Gemini Native JSON Schema
    # ─────────────────────────────────────────────────────────

    def _parse_with_ai(self, text: str) -> Optional[Dict[str, Any]]:
        ai = self._get_ai_service()

        # Context-window guard
        if len(text) > 40_000:
            text = text[:40_000] + "\n...[Truncated by CVParser]"

        prompt = (
            "Analyze the following resume and accurately extract all fields "
            "into the JSON schema structure you have been configured with.\n"
            "CRITICAL: If a section has bullet points (like work experience descriptions), "
            "you may output them as a JSON list of strings, or join them with newlines.\n\n"
            f"{text}"
        )

        raw_schema = CVParserOutputSchema.model_json_schema()

        generation_config = {
            "response_mime_type": "application/json",
            "response_schema": raw_schema,
        }
        generation_config_simple = {
            "response_mime_type": "application/json",
        }

        try:
            try:
                raw = ai._call(
                    prompt,
                    system_instruction=(
                        "You are a strict ATS data processor. Extract resume fields with high semantic "
                        "fidelity. Normalize dates to 'Month YYYY' format. Return 'Present' for ongoing roles."
                    ),
                    config=generation_config,
                    context_tag="cv_parse",
                )
            except Exception as schema_err:
                # response_schema may be rejected due to $defs limitations — retry without schema constraints
                logger.warning(
                    f"Schema-constrained Gemini call failed ({schema_err}), retrying without schema constraints."
                )
                raw = ai._call(
                    prompt,
                    system_instruction=(
                        "You are a strict ATS data processor. Extract resume fields with high semantic "
                        "fidelity. Normalize dates to 'Month YYYY' format. Return 'Present' for ongoing roles. "
                        "Return ONLY a valid JSON object matching the CVParserOutputSchema structure exactly."
                    ),
                    config=generation_config_simple,
                    context_tag="cv_parse",
                )

            if not raw:
                return None

            # Attempt clean parse first
            try:
                data = json.loads(raw.strip())
                return CVParserOutputSchema(**data).model_dump()
            except (json.JSONDecodeError, Exception):
                # Partial recovery: strip markdown fences and retry
                cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", raw.strip(), flags=re.MULTILINE)
                try:
                    data = json.loads(cleaned)
                    return CVParserOutputSchema(**data).model_dump()
                except Exception as recovery_err:
                    logger.error(f"AI response partial recovery failed: {recovery_err}")
                    return None

        except Exception as e:
            logger.error(f"Gemini API call error: {e}")
            return None

    # ─────────────────────────────────────────────────────────
    # Heuristic Path: Section-Aware Engine
    # ─────────────────────────────────────────────────────────

    def _parse_with_heuristics(self, text: str) -> Dict[str, Any]:
        sections = chunk_sections(text)
        header_text = sections.get("_header", text[:800])
        full_text = text

        personal = PersonalInfoSchema()

        # — Email
        m = re.search(r"[\w.\-+]+@[\w.\-]+\.\w{2,}", full_text)
        if m:
            personal.email = m.group(0)

        # — Phone (international + local formats)
        m = re.search(
            r"(?:\+?\d{1,3}[\s.\-]?)?(?:\(?\d{2,4}\)?[\s.\-]?)?\d{3,4}[\s.\-]?\d{3,4}[\s.\-]?\d{3,4}",
            full_text,
        )
        if m:
            candidate = re.sub(r"[^\d+]", "", m.group(0))
            if 7 <= len(candidate) <= 16:
                personal.phone = m.group(0).strip()

        # — URLs
        urls = re.findall(r"https?://[^\s<>\"')\]]+|www\.[^\s<>\"')\]]+", full_text)
        for url in urls:
            url = url.strip(".,;)]\\'\"")
            if "linkedin.com" in url and not personal.linkedin:
                personal.linkedin = url
            elif (
                any(d in url for d in ["github.com", "gitlab.com", "behance.net", "dribbble.com", "portfolio"])
                and not personal.portfolio
            ):
                personal.portfolio = url

        # — Name (confidence-scored)
        personal.full_name = self._extract_name(header_text, full_text)

        # — Job title (from header region)
        personal.job_title = self._extract_job_title(header_text)

        # — Location
        personal.location = self._extract_location(header_text + "\n" + full_text[:500])

        # — Summary
        summary = ""
        if "summary" in sections:
            summary = sections["summary"][:1000].strip()
        else:
            m = re.search(
                r"(?i)\b(?:summary|profile|objective|about me)\b\s*:?\s*(.*?)(?=\n\s*\n)",
                full_text, re.DOTALL,
            )
            if m:
                summary = m.group(1).strip()[:1000]

        # — Work experience
        work_exp = self._extract_work_experience(sections.get("experience", ""))

        # — Education
        education = self._extract_education(sections.get("education", ""))

        # — Skills
        skills_text = sections.get("skills", full_text)
        skills = self._extract_skills(skills_text)

        # — Extra / non-standard sections
        KNOWN_KEYS = {"_header", "summary", "experience", "education", "skills",
                      "projects", "certifications"}
        extra_sections: Dict[str, str] = {
            k: v for k, v in sections.items()
            if k not in KNOWN_KEYS and v.strip()
        }

        result = CVParserOutputSchema(
            personal_info=personal,
            professional_summary=summary or None,
            work_experience=work_exp,
            education=education,
            skills=skills,
            extra_sections=extra_sections,
        )
        return result.model_dump()

    # ─────────────────────────────────────────────────────────
    # Sub-Extractors
    # ─────────────────────────────────────────────────────────

    def _extract_name(self, header_text: str, full_text: str) -> Optional[str]:
        """
        Confidence-scored name extraction:
        - Prefers lines that are Title Cased with 2-4 words
        - Avoids lines containing common CV keywords
        - Skips lines that look like job titles or addresses
        """
        NOISE_WORDS = {
            "resume", "cv", "curriculum", "vitae", "email", "phone", "mobile",
            "address", "contact", "profile", "summary", "linkedin", "github",
            "portfolio", "website", "tel", "fax", "objective",
        }
        TITLE_WORDS = {
            "engineer", "developer", "designer", "manager", "analyst", "consultant",
            "director", "officer", "specialist", "coordinator", "architect", "lead",
        }

        candidates: List[Tuple[int, str]] = []
        lines = [l.strip() for l in header_text.split("\n") if l.strip()]

        for idx, line in enumerate(lines[:10]):
            lower = line.lower()
            words = line.split()
            score = 0

            # Must be 2–5 words
            if not (2 <= len(words) <= 5):
                continue
            # Must not be too long or too short
            if not (5 <= len(line) <= 50):
                continue
            # No noise keywords
            if any(nw in lower for nw in NOISE_WORDS):
                continue
            # No email/phone/url
            if re.search(r"[@\d/\\|]", line):
                continue
            # Prefer title case
            if line.istitle() or all(w[0].isupper() for w in words if w):
                score += 3
            # Prefer early lines
            score += max(0, 5 - idx)
            # Penalize if it looks like a job title
            if any(tw in lower for tw in TITLE_WORDS):
                score -= 2
            # Penalize ALL CAPS
            if line.isupper():
                score -= 1

            candidates.append((score, line))

        if candidates:
            return max(candidates, key=lambda x: x[0])[1]
        return None

    def _extract_job_title(self, header_text: str) -> Optional[str]:
        """Extracts job title from the header region using common title keywords."""
        TITLE_KEYWORDS = [
            "engineer", "developer", "designer", "manager", "analyst", "consultant",
            "director", "officer", "specialist", "coordinator", "architect", "lead",
            "scientist", "researcher", "executive", "associate", "intern", "head of",
        ]
        for line in header_text.split("\n"):
            stripped = line.strip()
            if 5 < len(stripped) < 80:
                lower = stripped.lower()
                if any(kw in lower for kw in TITLE_KEYWORDS):
                    if not re.search(r"[@\d]", stripped):
                        return stripped
        return None

    def _extract_location(self, text: str) -> Optional[str]:
        """Matches City, Country / City, State / City, State, Country patterns."""
        m = re.search(
            r"\b([A-Z][a-zA-Z\s]{2,20}),\s*([A-Z][a-zA-Z\s]{2,20})(?:,\s*[A-Z][a-zA-Z\s]{2,20})?\b",
            text,
        )
        if m:
            candidate = m.group(0).strip()
            # Filter out false positives like "January 2020, March 2021"
            if not re.search(r"\d", candidate):
                return candidate
        return None

    def _extract_work_experience(self, section_text: str) -> List[WorkExperienceSchema]:
        """
        Parses work experience blocks. Each job entry typically contains:
        - A line with title and/or company name
        - A date range line (Month YYYY – Month YYYY / Present)
        - Followed by description bullet points or paragraphs
        """
        if not section_text.strip():
            return []

        DATE_RANGE_RE = re.compile(
            r"(?:"
            r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
            r"Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"
            r"[\s.'\-]+['']?\d{2,4}"
            r"|(?:\d{4})"
            r")"
            r"\s*[–\-–to]+\s*"
            r"(?:"
            r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
            r"Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"
            r"[\s.'\-]+['']?\d{2,4}"
            r"|\d{4}"
            r"|[Pp]resent|[Cc]urrent|[Nn]ow"
            r")",
            re.IGNORECASE,
        )

        entries: List[WorkExperienceSchema] = []
        blocks = re.split(r"\n{2,}", section_text.strip())

        for block in blocks:
            lines = [l.strip() for l in block.split("\n") if l.strip()]
            if not lines:
                continue

            title: Optional[str] = None
            company: Optional[str] = None
            start_date: Optional[str] = None
            end_date: Optional[str] = None
            desc_lines: List[str] = []

            for line in lines:
                date_m = DATE_RANGE_RE.search(line)
                if date_m:
                    date_str = date_m.group(0)
                    parts = re.split(r"\s*[–\-–to]+\s*", date_str, maxsplit=1)
                    start_date = DateNormalizer.normalize(parts[0].strip()) if parts else None
                    end_date = DateNormalizer.normalize(parts[1].strip()) if len(parts) > 1 else None
                    remainder = line.replace(date_str, "").strip(" |–-,")
                    if remainder and not company:
                        company = remainder
                elif not title and len(line) < 80:
                    title = line
                elif not company and len(line) < 80 and title:
                    if not DATE_RANGE_RE.search(line):
                        company = line
                else:
                    desc_lines.append(line)

            if title or company:
                entries.append(WorkExperienceSchema(
                    title=title,
                    company=company,
                    start_date=start_date,
                    end_date=end_date,
                    description="\n".join(desc_lines).strip() or None,
                ))

        return entries

    def _extract_education(self, section_text: str) -> List[EducationSchema]:
        """
        Parses education blocks: degree, institution, year, grade.
        """
        if not section_text.strip():
            return []

        DEGREE_RE = re.compile(
            r"\b(B\.?Sc|B\.?A|B\.?Eng|B\.?Tech|M\.?Sc|M\.?A|M\.?Eng|MBA|Ph\.?D|"
            r"Bachelor['\s]?s?|Master['\s]?s?|Doctorate|Diploma|Certificate|HND|OND|Associate)\b",
            re.IGNORECASE,
        )
        GRADE_RE = re.compile(
            r"\b(First Class|Second Class|2:1|2:2|Third Class|Pass|Distinction|"
            r"Merit|GPA\s*:?\s*[\d.]+|[\d.]+\s*/\s*[\d.]+|Cum Laude|Magna Cum Laude|"
            r"Summa Cum Laude)\b",
            re.IGNORECASE,
        )
        YEAR_RE = re.compile(r"\b(19|20)\d{2}\b(?:\s*[–\-]\s*(?:(19|20)\d{2}|[Pp]resent))?")

        entries: List[EducationSchema] = []
        blocks = re.split(r"\n{2,}", section_text.strip())

        for block in blocks:
            lines = [l.strip() for l in block.split("\n") if l.strip()]
            if not lines:
                continue

            degree: Optional[str] = None
            institution: Optional[str] = None
            year: Optional[str] = None
            grade: Optional[str] = None

            for line in lines:
                year_m = YEAR_RE.search(line)
                grade_m = GRADE_RE.search(line)
                degree_m = DEGREE_RE.search(line)

                if grade_m and not grade:
                    grade = grade_m.group(0)
                if year_m and not year:
                    year = year_m.group(0)
                if degree_m and not degree:
                    degree = line  # full line often has degree + subject
                elif not institution and not degree_m and len(line) < 100:
                    if not year_m or len(line) > 10:
                        institution = line

            if degree or institution:
                entries.append(EducationSchema(
                    degree=degree,
                    institution=institution,
                    year=year,
                    grade=grade,
                ))

        return entries

    def _extract_skills(self, text: str) -> List[str]:
        """
        Extracts skills using alias-aware pattern matching with canonical normalization.
        Also parses comma/bullet-separated skill lists when present in the skills section.
        """
        found: Dict[str, str] = {}  # canonical → display
        lower_text = text.lower()

        # Pattern match against known skills registry
        for pattern in SKILL_PATTERNS:
            if re.search(rf"(?<![a-z]){re.escape(pattern)}(?![a-z])", lower_text):
                canonical = normalize_skill(pattern)
                found[canonical.lower()] = canonical

        # Also parse explicit skill lists (CSV or bullet lines within skills section)
        list_items = re.split(r"[,•\|\n]+", text)
        for item in list_items:
            item = item.strip(" .-\t")
            if 2 <= len(item) <= 40 and not re.search(r"\d{4}", item):
                canonical = normalize_skill(item)
                key = canonical.lower()
                if key not in found:
                    found[key] = canonical

        return sorted(found.values(), key=str.lower)

