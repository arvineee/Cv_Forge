"""
DOCX Generation Service — Fixed Edition

Fixes applied in this pass:
- _normalize() now carries EVERY resume field (certifications, achievements,
  languages, interests, awards, publications, volunteer, references, projects,
  skill_groups, github). Previously only 6 fields were read, so every DOCX
  download silently dropped those sections regardless of what was in the DB.
- _add_job() now reads job_title (matching the schema used everywhere else in
  the app) instead of the nonexistent "title" key, and renders the
  achievements[] bullet list, not just a free-text description.
- generate() now pulls custom_sections / certifications / languages / awards /
  references / projects the same way PDFService.generate() already did, so
  DOCX and PDF exports are at parity.
- Skill groups (with subtitles) are rendered when present, falling back to a
  flat comma list otherwise.
- github is read from personal_info.github separately instead of being
  guessed out of the portfolio URL.

Improvements kept from v1:
- generate_from_dict() works with plain dicts (no ORM model needed).
- Style setup is document-scoped — never mutates shared global style objects.
- Section headings use paragraph bottom-border dividers (not heading style mutation).
- Title / date on job entries use tab stops for true left+right alignment.
- Bullet points use List Bullet numbering style (no manual unicode bullet chars).
- LinkedIn / portfolio / GitHub rendered as live hyperlinks via oxml relationship injection.
- Accent color is a single constant — change one value to retheme the whole doc.
- Page number injected into the footer section.
- All fields fail gracefully — None values are skipped, never rendered as "None".
"""

import os
import re
import tempfile
import logging
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────────────────────
# Theme / Constants
# ─────────────────────────────────────────────────────────────

ACCENT_RGB    = RGBColor(0x25, 0x63, 0xEB)   # Blue-600
HEADING_RGB   = RGBColor(0x0F, 0x17, 0x2A)   # Slate-900
MUTED_RGB     = RGBColor(0x47, 0x55, 0x69)   # Slate-600
DIVIDER_HEX   = "CBD5E1"                      # Slate-300 (no #)

CONTENT_WIDTH_INCHES = 6.5
RIGHT_TAB_TWIPS      = int(CONTENT_WIDTH_INCHES * 1440)  # twips


# ─────────────────────────────────────────────────────────────
# Resume Data Normalizer
# ─────────────────────────────────────────────────────────────

def _normalize(resume: Any) -> Dict[str, Any]:
    """
    Accept ORM model or plain dict; return a clean, safe dict with EVERY
    section the app supports. Mirrors pdf_service._normalize_data() so
    DOCX and PDF exports never drift apart again.
    """
    if isinstance(resume, dict):
        data = dict(resume)
    else:
        custom = getattr(resume, "custom_sections", None) or {}
        data = {
            "personal_info":        getattr(resume, "personal_info", {}) or {},
            "professional_summary": getattr(resume, "professional_summary", None),
            "work_experience":      getattr(resume, "work_experience", []) or [],
            "education":            getattr(resume, "education", []) or [],
            "skills":               getattr(resume, "skills", []) or [],
            "skill_groups":         custom.get("skill_groups", []),
            "certifications":       getattr(resume, "certifications", []) or [],
            "achievements":         custom.get("achievements", []),
            "languages":            getattr(resume, "languages", []) or [],
            "interests":            custom.get("interests", []),
            "awards":               getattr(resume, "awards", []) or [],
            "publications":         custom.get("publications", []),
            "volunteer":            custom.get("volunteer", ""),
            "references":           getattr(resume, "references", None),
            "projects":             getattr(resume, "projects", []) or [],
            "title":                getattr(resume, "title", None),
        }

    personal = data.get("personal_info") or {}
    if not isinstance(personal, dict):
        personal = {}

    name = (
        personal.get("full_name")
        or f"{personal.get('first_name', '')} {personal.get('last_name', '')}".strip()
        or data.get("title")
        or "Resume"
    )

    raw_work = data.get("work_experience") or []
    work = []
    for job in raw_work:
        if not isinstance(job, dict):
            continue
        work.append({
            "job_title":    job.get("job_title") or job.get("title") or "",
            "company":      job.get("company") or "",
            "location":     job.get("location") or "",
            "start_date":   job.get("start_date") or "",
            "end_date":     job.get("end_date") or "Present",
            "description":  job.get("description") or "",
            "achievements": job.get("achievements") or [],
        })

    raw_edu = data.get("education") or []
    education = []
    for e in raw_edu:
        if not isinstance(e, dict):
            continue
        education.append({
            "degree":      e.get("degree") or "",
            "institution": e.get("institution") or "",
            "location":    e.get("location") or "",
            "year":        e.get("year") or "",
            "grade":       e.get("grade") or "",
        })

    raw_projects = data.get("projects") or []
    projects = []
    for p in raw_projects:
        if isinstance(p, dict):
            projects.append({
                "name": p.get("name") or "",
                "description": p.get("description") or "",
                "url": p.get("url") or "",
            })
        elif isinstance(p, str) and p.strip():
            projects.append({"name": p.strip(), "description": "", "url": ""})

    return {
        "name":           name,
        "job_title":      personal.get("job_title") or "",
        "email":          personal.get("email") or "",
        "phone":          personal.get("phone") or "",
        "location":       personal.get("location") or "",
        "linkedin":       personal.get("linkedin") or "",
        "portfolio":      personal.get("portfolio") or personal.get("website") or "",
        "github":         personal.get("github") or "",
        "summary":        data.get("professional_summary") or data.get("objective") or "",
        "work":           work,
        "education":      education,
        "skills":         data.get("skills") or [],
        "skill_groups":   data.get("skill_groups") or [],
        "certifications": data.get("certifications") or [],
        "achievements":   data.get("achievements") or [],
        "languages":      data.get("languages") or [],
        "interests":      data.get("interests") or [],
        "awards":         data.get("awards") or [],
        "publications":   data.get("publications") or [],
        "volunteer":      data.get("volunteer") or "",
        "references":     data.get("references") or "",
        "projects":       projects,
    }


# ─────────────────────────────────────────────────────────────
# oxml Helpers
# ─────────────────────────────────────────────────────────────

def _set_paragraph_bottom_border(para, color_hex: str = DIVIDER_HEX, size: int = 6):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_hyperlink(para, text: str, url: str, rgb: RGBColor = ACCENT_RGB):
    part = para.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_elem = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    color_elem = OxmlElement("w:color")
    color_elem.set(qn("w:val"), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    rPr.append(color_elem)

    run_elem.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run_elem.append(t)
    hyperlink.append(run_elem)
    para._p.append(hyperlink)


def _add_tab_stop(para, position_twips: int, alignment=WD_TAB_ALIGNMENT.RIGHT):
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right" if alignment == WD_TAB_ALIGNMENT.RIGHT else "left")
    tab.set(qn("w:pos"), str(position_twips))
    tabs.append(tab)
    pPr.append(tabs)


def _add_page_number_field(para):
    run = para.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar_begin)

    instrText = OxmlElement("w:instrText")
    instrText.text = " PAGE "
    run._r.append(instrText)

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar_end)


# ─────────────────────────────────────────────────────────────
# Main Service
# ─────────────────────────────────────────────────────────────

class DOCXService:
    """
    Generate a polished DOCX resume.

    Usage:
        path = DOCXService().generate(resume_obj)
        path = DOCXService().generate_from_dict(data_dict)
    """

    def generate(self, resume) -> str:
        return self._build_docx(_normalize(resume))

    def generate_from_dict(self, data: Dict[str, Any]) -> str:
        return self._build_docx(_normalize(data))

    def _build_docx(self, d: Dict[str, Any]) -> str:
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)

        doc = Document()
        self._setup_styles(doc)
        self._setup_page(doc)

        self._add_name(doc, d["name"])

        if d["job_title"]:
            self._add_centered_run(doc, d["job_title"], size=12, rgb=ACCENT_RGB)

        contact_parts = [v for v in [d["email"], d["phone"], d["location"]] if v]
        if contact_parts:
            self._add_centered_run(doc, "  |  ".join(contact_parts), size=9, rgb=MUTED_RGB)

        link_items = []
        if d["linkedin"]:
            link_items.append(("LinkedIn", d["linkedin"]))
        if d["github"]:
            link_items.append(("GitHub", d["github"]))
        if d["portfolio"]:
            link_items.append(("Portfolio", d["portfolio"]))
        if link_items:
            link_para = doc.add_paragraph()
            link_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for i, (label, url) in enumerate(link_items):
                if i > 0:
                    link_para.add_run("  ·  ").font.color.rgb = MUTED_RGB
                url = url if url.startswith("http") else f"https://{url}"
                _add_hyperlink(link_para, label, url, ACCENT_RGB)

        doc.add_paragraph()

        if d["summary"]:
            self._section_heading(doc, "PROFESSIONAL SUMMARY")
            doc.add_paragraph(d["summary"])

        if d["work"]:
            self._section_heading(doc, "WORK EXPERIENCE")
            for job in d["work"]:
                self._add_job(doc, job)

        if d["education"]:
            self._section_heading(doc, "EDUCATION")
            for edu in d["education"]:
                self._add_education(doc, edu)

        if d["skill_groups"] or d["skills"]:
            self._section_heading(doc, "SKILLS")
            self._add_skills(doc, d["skill_groups"], d["skills"])

        if d["certifications"]:
            self._section_heading(doc, "CERTIFICATIONS")
            self._add_bullet_list(doc, d["certifications"])

        if d["achievements"]:
            self._section_heading(doc, "KEY ACHIEVEMENTS")
            self._add_bullet_list(doc, d["achievements"])

        if d["awards"]:
            self._section_heading(doc, "AWARDS")
            self._add_bullet_list(doc, d["awards"])

        if d["projects"]:
            self._section_heading(doc, "PROJECTS")
            for proj in d["projects"]:
                self._add_project(doc, proj)

        if d["publications"]:
            self._section_heading(doc, "PUBLICATIONS")
            self._add_bullet_list(doc, d["publications"])

        if d["volunteer"]:
            self._section_heading(doc, "VOLUNTEER EXPERIENCE")
            doc.add_paragraph(d["volunteer"])

        if d["languages"]:
            self._section_heading(doc, "LANGUAGES")
            p = doc.add_paragraph(" · ".join(str(x) for x in d["languages"]))
            for run in p.runs:
                run.font.name = "Calibri"

        if d["interests"]:
            self._section_heading(doc, "INTERESTS")
            p = doc.add_paragraph(" · ".join(str(x) for x in d["interests"]))
            for run in p.runs:
                run.font.name = "Calibri"

        if d["references"]:
            self._section_heading(doc, "REFERENCES")
            doc.add_paragraph(d["references"])

        self._add_footer(doc, d["name"])

        doc.save(path)
        logger.info(f"DOCX generated at: {path}")
        return path

    def _setup_page(self, doc: Document):
        section = doc.sections[0]
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
        margin = Inches(0.75)
        section.top_margin    = margin
        section.bottom_margin = Inches(0.65)
        section.left_margin   = margin
        section.right_margin  = margin

    def _setup_styles(self, doc: Document):
        try:
            normal = doc.styles["Normal"]
            normal.font.name = "Calibri"
            normal.font.size = Pt(10)

            for level in range(1, 4):
                h = doc.styles[f"Heading {level}"]
                h.font.name = "Calibri"
                h.font.bold = True
                h.font.color.rgb = HEADING_RGB
                h.paragraph_format.space_before = Pt(0)
                h.paragraph_format.space_after  = Pt(2)

            if "List Bullet" in doc.styles:
                lb = doc.styles["List Bullet"]
                lb.font.name = "Calibri"
                lb.font.size = Pt(10)
        except Exception as e:
            logger.warning(f"Style configuration warning: {e}")

    def _add_name(self, doc: Document, name: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        run.bold = True
        run.font.size = Pt(26)
        run.font.name = "Calibri"
        run.font.color.rgb = HEADING_RGB

    def _add_centered_run(self, doc: Document, text: str, size: int = 10, rgb: RGBColor = HEADING_RGB):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.color.rgb = rgb
        run.font.name = "Calibri"

    def _section_heading(self, doc: Document, title: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
        _set_paragraph_bottom_border(p, DIVIDER_HEX, size=6)

        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run.font.color.rgb = ACCENT_RGB

    def _add_job(self, doc: Document, job: Dict[str, Any]):
        """
        Renders one work experience entry. Fixed: reads job_title (not the
        nonexistent "title" key), and renders the achievements[] bullet
        list in addition to any free-text description.
        """
        title     = job.get("job_title") or job.get("title") or ""
        company   = job.get("company") or ""
        location  = job.get("location") or ""
        start     = job.get("start_date") or ""
        end       = job.get("end_date") or "Present"
        date_str  = f"{start} – {end}" if start else end
        desc      = job.get("description") or ""
        achievements = job.get("achievements") or []

        if title or date_str:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            _add_tab_stop(p, RIGHT_TAB_TWIPS)

            if title:
                run = p.add_run(title)
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = "Calibri"
                run.font.color.rgb = HEADING_RGB

            if date_str:
                tab_run = p.add_run(f"\t{date_str}")
                tab_run.font.size = Pt(9)
                tab_run.font.color.rgb = MUTED_RGB
                tab_run.font.name = "Calibri"

        company_loc = ", ".join(v for v in [company, location] if v)
        if company_loc:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(company_loc)
            run.font.size = Pt(9.5)
            run.font.color.rgb = MUTED_RGB
            run.font.name = "Calibri"

        if desc:
            self._render_description(doc, desc)

        for a in achievements:
            if isinstance(a, str) and a.strip():
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(a.strip())
                run.font.name = "Calibri"
                run.font.size = Pt(10)

        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def _add_education(self, doc: Document, edu: Dict[str, Any]):
        degree      = edu.get("degree") or ""
        institution = edu.get("institution") or ""
        year        = edu.get("year") or ""
        grade       = edu.get("grade") or ""

        if degree:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)

            run = p.add_run(degree)
            run.bold = True
            run.font.name = "Calibri"
            run.font.color.rgb = HEADING_RGB

            if year:
                yr_run = p.add_run(f"  ({year})")
                yr_run.font.size = Pt(9.5)
                yr_run.font.color.rgb = MUTED_RGB
                yr_run.font.name = "Calibri"

        meta_parts = [v for v in [institution, grade] if v]
        if meta_parts:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run("  ·  ".join(meta_parts))
            run.font.size = Pt(9.5)
            run.font.color.rgb = MUTED_RGB
            run.font.name = "Calibri"

    def _add_skills(self, doc: Document, skill_groups: List[Dict], flat_skills: List[str]):
        """Render grouped skills with subtitles when available, else a flat comma list."""
        if skill_groups:
            for group in skill_groups:
                group_name = group.get("group") if isinstance(group, dict) else getattr(group, "group", None)
                skills = group.get("skills") if isinstance(group, dict) else getattr(group, "skills", [])
                if not skills:
                    continue
                if group_name:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(1)
                    run = p.add_run(group_name)
                    run.bold = True
                    run.font.size = Pt(9.5)
                    run.font.name = "Calibri"
                    run.font.color.rgb = MUTED_RGB
                doc.add_paragraph(", ".join(str(s) for s in skills))
        elif flat_skills:
            doc.add_paragraph(", ".join(str(s) for s in flat_skills))

    def _add_bullet_list(self, doc: Document, items: List[str]):
        """Generic bullet-list renderer used by certifications/achievements/awards/publications."""
        for item in items:
            if not item:
                continue
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(str(item).strip())
            run.font.name = "Calibri"
            run.font.size = Pt(10)

    def _add_project(self, doc: Document, proj: Dict[str, Any]):
        name = proj.get("name") or ""
        description = proj.get("description") or ""
        url = proj.get("url") or ""
        if not (name or description):
            return
        if name:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(name)
            run.bold = True
            run.font.name = "Calibri"
            run.font.color.rgb = HEADING_RGB
        if url:
            link_p = doc.add_paragraph()
            link_p.paragraph_format.space_after = Pt(1)
            _add_hyperlink(link_p, url, url if url.startswith("http") else f"https://{url}", ACCENT_RGB)
        if description:
            self._render_description(doc, description)
        doc.add_paragraph().paragraph_format.space_after = Pt(1)

    def _render_description(self, doc: Document, text: str):
        BULLET_RE = re.compile(r"^[\•\-\*\–\u2022]\s*|^\d+\.\s+")
        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            if BULLET_RE.match(stripped):
                clean = BULLET_RE.sub("", stripped).strip()
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(clean)
                run.font.name = "Calibri"
                run.font.size = Pt(10)
            else:
                p = doc.add_paragraph(stripped)
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.name = "Calibri"

    def _add_footer(self, doc: Document, candidate_name: str):
        try:
            section = doc.sections[0]
            footer = section.footer
            footer.is_linked_to_previous = False

            para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            para.clear()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            name_run = para.add_run(f"{candidate_name}  ·  Page ")
            name_run.font.size = Pt(8)
            name_run.font.color.rgb = MUTED_RGB
            name_run.font.name = "Calibri"

            _add_page_number_field(para)

        except Exception as e:
            logger.warning(f"Footer injection failed: {e}")

